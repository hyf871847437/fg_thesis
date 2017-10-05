# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches

def write_thesis(path,title,text):

    document = Document()

    document.add_heading('%s'%title, 0)

    # document.add_paragraph('\n'*10)
    # document.add_paragraph('\t'*2+ u'考查科目 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'学    院 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'专    业 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'年    级 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'学    期 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'姓    名 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'学    号 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\t'*2+ u'成    绩 ————————————————', style='IntenseQuote')
    # document.add_paragraph('\n' * 10)
    # #document.add_paragraph('first item in unordered list', style='ListBullet')
    # p = document.add_paragraph(u'【摘要】A plain paragraph having some')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    # p.add_run(u'【自己写】').bold = True
    # q = document.add_paragraph(u'【关键词】A B C D')
    # q.add_run(u'【自己写】').bold = True
    #
    # document.add_heading(u'目录', level=2)
    # document.add_paragraph('first item in ordered list', style='ListNumber')
    # document.add_paragraph(' ....', style='ListNumber')
    # document.add_heading(u'正文', level=2)

    s = document.add_paragraph(text)
    # document.add_picture('monty-truth.png', width=Inches(1.25))
    # table = document.add_table(rows=1, cols=3)

    document.add_page_break()

    document.save(path)
    print 'over'

